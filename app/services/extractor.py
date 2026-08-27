"""
CV text extraction from PDF and Word-family files (.docx / .doc / .rtf).
Returns plain text suitable for passing to the Claude API.

Extraction strategy:
  - PDF            → pdfminer directly.
  - Word-family    → convert to PDF with headless LibreOffice, then pdfminer.
                     This captures text that python-docx silently drops (Word
                     text boxes, drawings, legacy binary .doc) and de-duplicates
                     Word's AlternateContent (modern + legacy copies). Falls back
                     to python-docx if LibreOffice is unavailable or fails.

Note: the builder still writes the OUTPUT .docx exactly as designed — LibreOffice
is used ONLY to read messy INPUT files, never to rewrite our output.
"""

import base64
import io
import logging
import os
import shutil
import subprocess
import tempfile

logger = logging.getLogger(__name__)

# LibreOffice binary (installed via the Dockerfile as `libreoffice-writer`).
_SOFFICE = shutil.which("soffice") or shutil.which("libreoffice")
_LO_TIMEOUT_SECONDS = 90


def extract_text(content_b64: str, content_type: str, filename: str) -> str:
    """
    Extract plain text from a base64-encoded CV file.

    Args:
        content_b64: Base64-encoded file content from Postmark attachment
        content_type: MIME type from Postmark (e.g. 'application/pdf')
        filename: Original filename — used as fallback to determine type

    Returns:
        Plain text content of the CV

    Raises:
        ValueError: If file type is unsupported or extraction fails
    """
    file_bytes = base64.b64decode(content_b64)
    ct = content_type.lower()
    fn = filename.lower()

    if "pdf" in ct or fn.endswith(".pdf"):
        return _extract_from_pdf(file_bytes)
    elif (
        "word" in ct
        or "openxmlformats" in ct
        or "opendocument" in ct
        or "rtf" in ct
        or fn.endswith((".docx", ".doc", ".rtf", ".odt"))
    ):
        return _extract_from_word(file_bytes, fn)
    else:
        raise ValueError(f"Unsupported file type: {content_type} / {filename}")


def _extract_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using pdfminer."""
    try:
        from pdfminer.high_level import extract_text_to_fp
        from pdfminer.layout import LAParams

        output = io.StringIO()
        extract_text_to_fp(
            io.BytesIO(file_bytes),
            output,
            laparams=LAParams(),
            output_type="text",
            codec="utf-8",
        )
        text = output.getvalue().strip()

        if text:
            return text

        # No text layer → likely a scanned / photographed / image-exported PDF.
        # Try OCR before giving up.
        from app.config import get_settings
        if get_settings().ocr_enabled:
            logger.info("PDF has no text layer; attempting OCR fallback.")
            ocr_text = _ocr_pdf(file_bytes)
            if ocr_text:
                logger.info(f"OCR fallback recovered {len(ocr_text)} chars.")
                return ocr_text

        raise ValueError("PDF appears to be scanned/image-only — no extractable text found.")

    except ImportError:
        raise RuntimeError("pdfminer.six is not installed. Run: pip install pdfminer.six")
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        raise ValueError(f"Could not extract text from PDF: {e}") from e


def _ocr_pdf(file_bytes: bytes) -> str:
    """
    OCR fallback for PDFs with no text layer (scanned / photographed / image-only).
    Rasterises each page with pdf2image (poppler) and runs Tesseract via
    pytesseract. Bounded by ocr_max_pages so a huge file can't stall the worker.
    Returns extracted text, or "" if OCR is unavailable or finds nothing.
    """
    try:
        import pytesseract
        from pdf2image import convert_from_bytes
    except ImportError:
        logger.warning("OCR libraries not installed (pytesseract / pdf2image); skipping OCR.")
        return ""

    from app.config import get_settings
    settings = get_settings()

    try:
        images = convert_from_bytes(
            file_bytes,
            dpi=settings.ocr_dpi,
            fmt="png",
            last_page=settings.ocr_max_pages,
        )
    except Exception as e:
        logger.error(f"OCR rasterisation failed: {e}")
        return ""

    pages = []
    for i, img in enumerate(images):
        try:
            pages.append(pytesseract.image_to_string(img, lang="eng"))
        except Exception as e:
            logger.error(f"OCR failed on page {i + 1}: {e}")

    return "\n".join(pages).strip()


def _extract_from_word(file_bytes: bytes, filename: str) -> str:
    """
    Word-family (.docx / .doc / .rtf): LibreOffice → PDF → pdfminer, with a
    python-docx fallback so the pipeline keeps working if LibreOffice is missing.
    """
    if _SOFFICE:
        try:
            text = _extract_via_libreoffice(file_bytes, filename)
            if text:
                return text
            logger.warning("LibreOffice produced no text; falling back to python-docx.")
        except Exception as e:
            logger.warning(f"LibreOffice extraction failed ({e}); falling back to python-docx.")
    else:
        logger.warning("LibreOffice (soffice) not found on PATH; using python-docx fallback.")

    return _extract_from_docx_pythondocx(file_bytes)


def _extract_via_libreoffice(file_bytes: bytes, filename: str) -> str:
    """
    Convert a Word-family file to PDF with headless LibreOffice, then run the
    existing pdfminer extractor. Uses a throwaway temp dir + per-run user profile
    so it is safe in a container, and enforces a timeout so a hung conversion
    can't stall the worker.
    """
    ext = os.path.splitext(filename)[1].lower()
    if ext not in (".docx", ".doc", ".rtf", ".odt"):
        ext = ".docx"

    with tempfile.TemporaryDirectory() as tmp:
        src = os.path.join(tmp, "input" + ext)
        with open(src, "wb") as fh:
            fh.write(file_bytes)

        profile = os.path.join(tmp, "lo_profile")
        cmd = [
            _SOFFICE,
            "--headless",
            "--norestore",
            "--nolockcheck",
            "--nodefault",
            f"-env:UserInstallation=file://{profile}",
            "--convert-to",
            "pdf",
            "--outdir",
            tmp,
            src,
        ]
        try:
            proc = subprocess.run(
                cmd,
                capture_output=True,
                timeout=_LO_TIMEOUT_SECONDS,
                env={**os.environ, "HOME": tmp},
            )
        except subprocess.TimeoutExpired as e:
            raise ValueError(
                f"LibreOffice conversion timed out after {_LO_TIMEOUT_SECONDS}s"
            ) from e

        pdf_path = os.path.join(tmp, "input.pdf")
        if proc.returncode != 0 or not os.path.exists(pdf_path):
            stderr = proc.stderr.decode("utf-8", "ignore")[:300]
            raise ValueError(
                f"LibreOffice conversion failed (rc={proc.returncode}): {stderr}"
            )

        with open(pdf_path, "rb") as fh:
            pdf_bytes = fh.read()

    return _extract_from_pdf(pdf_bytes)


def _extract_from_docx_pythondocx(file_bytes: bytes) -> str:
    """Fallback: extract text from DOCX using python-docx (paragraphs + tables).

    Note: python-docx does NOT read Word text boxes/drawings, so this is a
    best-effort fallback only — the LibreOffice path above is preferred.
    """
    try:
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        paragraphs = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_texts:
                    paragraphs.append(" | ".join(row_texts))

        text = "\n".join(paragraphs).strip()

        if not text:
            raise ValueError(
                "DOCX appears to be empty — no extractable text found "
                "(content may be in text boxes that require LibreOffice)."
            )

        return text

    except ImportError:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx")
    except Exception as e:
        logger.error(f"DOCX (python-docx) extraction failed: {e}")
        raise ValueError(f"Could not extract text from DOCX: {e}") from e
