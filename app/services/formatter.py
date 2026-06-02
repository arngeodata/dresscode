"""
CV formatter — style-guide-driven DOCX builder.

Takes a ParsedCV and a style guide dict, builds a clean DOCX from scratch
using python-docx. No placeholder templates — the style guide drives fonts,
colours, and layout. Sensible defaults apply if any field is missing.

Optional header_image_bytes: if supplied, the image is embedded into the Word
header section (full-width, flush to the top of the page). Used when an org's
style_guide specifies header.image_path.
"""

import io
import logging
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.models import ParsedCV

logger = logging.getLogger(__name__)

# ── Default style guide (fallback for any missing keys) ───────────────────────
DEFAULT_STYLE_GUIDE = {
    "fonts": {
        "name_font":    "Calibri",
        "body_font":    "Calibri",
        "name_size":    22,
        "section_size": 11,
        "body_size":    10,
    },
    "colours": {
        "primary_hex": "1B3A6B",
        "accent_hex":  "4472C4",
        "text_hex":    "000000",
        "contact_hex": "595959",
    },
    "layout": {
        "margins_cm":     1.8,
        "name_alignment": "left",   # left | center
        "section_border": True,     # coloured rule under section headings
        "name_suffix":    "",       # e.g. " – CV" for Hyperion style
    },
    "sections": {
        "order":            ["summary", "experience", "education", "skills"],
        "summary_label":    "Profile",
        "experience_label": "Experience",
        "education_label":  "Education",
        "skills_label":     "Key Skills",
    },
    "header": {
        "contact_separator":       "  |  ",
        "show_linkedin":           True,
        "suppress_contact_details": False,  # set True to omit contact line entirely
        # image_bucket / image_path: handled by worker.py — not used here directly
    },
    "output": {
        # filename_format: use {name} placeholder; omit to keep default naming
        "filename_format": "",
    },
}


def _merge(defaults: dict, overrides: dict) -> dict:
    """Deep-merge overrides into defaults."""
    result = defaults.copy()
    for key, val in overrides.items():
        if isinstance(val, dict) and isinstance(result.get(key), dict):
            result[key] = _merge(result[key], val)
        else:
            result[key] = val
    return result


def _hex_to_rgb(hex_str: str) -> RGBColor:
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _add_rule(paragraph, colour_hex: str):
    """Add a thin coloured bottom border to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), colour_hex.lstrip("#"))
    pBdr.append(bottom)
    pPr.append(pBdr)


def build_cv_docx(
    cv: ParsedCV,
    style_guide: dict | None = None,
    header_image_bytes: bytes | None = None,
) -> bytes:
    """
    Build a formatted DOCX from a ParsedCV and an optional style guide dict.

    Args:
        cv:                  Structured CV data from Claude parsing.
        style_guide:         Style parameters. Missing keys fall back to
                             DEFAULT_STYLE_GUIDE. Pass None to use full defaults.
        header_image_bytes:  Raw PNG/JPEG bytes for the branded header image.
                             If supplied, the image is embedded full-width into
                             the Word header section, flush to the top of the page.

    Returns:
        Raw DOCX bytes ready for email attachment.
    """
    sg = _merge(DEFAULT_STYLE_GUIDE, style_guide or {})

    fonts    = sg["fonts"]
    colours  = sg["colours"]
    layout   = sg["layout"]
    sections = sg["sections"]
    header   = sg["header"]

    name_font    = fonts["name_font"]
    body_font    = fonts["body_font"]
    name_size    = fonts["name_size"]
    section_size = fonts["section_size"]
    body_size    = fonts["body_size"]

    use_section_border       = layout.get("section_border", True)
    name_suffix              = layout.get("name_suffix", "")
    suppress_contact_details = header.get("suppress_contact_details", False)

    primary_colour = _hex_to_rgb(colours["primary_hex"])
    accent_colour  = _hex_to_rgb(colours["accent_hex"])
    text_colour    = _hex_to_rgb(colours["text_hex"])
    contact_colour = _hex_to_rgb(colours.get("contact_hex", colours["text_hex"]))

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for sec in doc.sections:
        m = Cm(layout["margins_cm"])
        sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = m

    # ── Branded header image ───────────────────────────────────────────────────
    if header_image_bytes:
        section = doc.sections[0]

        # Set header-from-top distance to 0 so image is flush to page edge
        pg_mar = section._sectPr.find(qn("w:pgMar"))
        if pg_mar is not None:
            pg_mar.set(qn("w:header"), "0")

        hdr = section.header
        hdr.is_linked_to_previous = False

        # Remove default blank paragraph from header
        for p in list(hdr.paragraphs):
            p._element.getparent().remove(p._element)

        # Add image paragraph — negative left indent extends to page edge
        hdr_para = hdr.add_paragraph()
        hdr_para.paragraph_format.space_before = Pt(0)
        hdr_para.paragraph_format.space_after  = Pt(0)
        hdr_para.paragraph_format.left_indent  = -Cm(layout["margins_cm"])
        img_run = hdr_para.add_run()
        img_run.add_picture(io.BytesIO(header_image_bytes), width=Inches(8.27))

        logger.info("Branded header image embedded")

    # Remove the default empty paragraph Word adds to the body
    for para in list(doc.paragraphs):
        para._element.getparent().remove(para._element)

    # ── Helpers ───────────────────────────────────────────────────────────────
    def add_run(paragraph, text, font=None, size=None, bold=False,
                italic=False, colour=None):
        run = paragraph.add_run(text)
        run.font.name      = font or body_font
        run.font.size      = Pt(size or body_size)
        run.bold           = bold
        run.italic         = italic
        run.font.color.rgb = colour or text_colour
        return run

    def add_section_header(label: str):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after  = Pt(4)
        run = para.add_run(label)
        run.bold           = True
        run.font.name      = body_font
        run.font.size      = Pt(section_size)
        run.font.color.rgb = primary_colour
        if use_section_border:
            _add_rule(para, colours["primary_hex"])
        return para

    def add_body(text, bold=False, italic=False, colour=None,
                 space_before=0, space_after=2):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(space_before)
        para.paragraph_format.space_after  = Pt(space_after)
        run = para.add_run(text or "")
        run.font.name      = body_font
        run.font.size      = Pt(body_size)
        run.bold           = bold
        run.italic         = italic
        run.font.color.rgb = colour or text_colour
        return para

    # ── Name alignment ────────────────────────────────────────────────────────
    name_align = (WD_ALIGN_PARAGRAPH.CENTER
                  if layout.get("name_alignment") == "center"
                  else WD_ALIGN_PARAGRAPH.LEFT)

    # ── Candidate name / document title ───────────────────────────────────────
    name_text = (cv.candidate.full_name or "") + name_suffix
    name_para = doc.add_paragraph()
    name_para.alignment = name_align
    name_para.paragraph_format.space_before = Pt(0)
    name_para.paragraph_format.space_after  = Pt(2)
    add_run(name_para, name_text,
            font=name_font, size=name_size, bold=True, colour=primary_colour)

    # ── Contact line (omitted when suppress_contact_details is True) ──────────
    if not suppress_contact_details:
        sep = header.get("contact_separator", "  |  ")
        contact_parts = []
        if cv.candidate.email:    contact_parts.append(cv.candidate.email)
        if cv.candidate.phone:    contact_parts.append(cv.candidate.phone)
        if cv.candidate.location: contact_parts.append(cv.candidate.location)
        if header.get("show_linkedin") and cv.candidate.linkedin:
            contact_parts.append(cv.candidate.linkedin)

        if contact_parts:
            c_para = doc.add_paragraph()
            c_para.alignment = name_align
            c_para.paragraph_format.space_before = Pt(0)
            c_para.paragraph_format.space_after  = Pt(6)
            add_run(c_para, sep.join(contact_parts),
                    size=body_size - 0.5, colour=contact_colour)

    # ── Sections ──────────────────────────────────────────────────────────────
    for section_key in sections.get("order", ["summary", "experience", "education", "skills"]):

        if section_key == "summary" and cv.summary:
            add_section_header(sections.get("summary_label", "Profile"))
            add_body(cv.summary, space_after=4)

        elif section_key == "experience" and cv.experience:
            add_section_header(sections.get("experience_label", "Experience"))

            for job in cv.experience:
                job_para = doc.add_paragraph()
                job_para.paragraph_format.space_before = Pt(6)
                job_para.paragraph_format.space_after  = Pt(1)
                add_run(job_para, job.title or "", bold=True)
                if job.company:
                    add_run(job_para, f"  |  {job.company}", colour=accent_colour)

                # Date range
                dates = []
                if job.start_date: dates.append(job.start_date)
                if job.end_date:   dates.append(job.end_date)
                elif job.start_date: dates.append("Present")
                if dates:
                    d_para = doc.add_paragraph()
                    d_para.paragraph_format.space_before = Pt(0)
                    d_para.paragraph_format.space_after  = Pt(2)
                    add_run(d_para, " – ".join(dates),
                            italic=True, size=body_size - 0.5, colour=contact_colour)

                # Bullets
                for bullet in (job.responsibilities or []):
                    bullet = (bullet or "").strip()
                    if not bullet:
                        continue
                    if bullet[:2] in ("- ", "• ", "* "):
                        bullet = bullet[2:]
                    bp = doc.add_paragraph(style="List Bullet")
                    bp.paragraph_format.space_before = Pt(0)
                    bp.paragraph_format.space_after  = Pt(1)
                    r = bp.add_run(bullet)
                    r.font.name      = body_font
                    r.font.size      = Pt(body_size)
                    r.font.color.rgb = text_colour

        elif section_key == "education" and cv.education:
            add_section_header(sections.get("education_label", "Education"))

            for edu in cv.education:
                edu_para = doc.add_paragraph()
                edu_para.paragraph_format.space_before = Pt(6)
                edu_para.paragraph_format.space_after  = Pt(1)
                add_run(edu_para, edu.qualification or "", bold=True)
                if edu.institution:
                    add_run(edu_para, f"  |  {edu.institution}", colour=accent_colour)

                if edu.year:
                    y_para = doc.add_paragraph()
                    y_para.paragraph_format.space_before = Pt(0)
                    y_para.paragraph_format.space_after  = Pt(2)
                    add_run(y_para, str(edu.year),
                            italic=True, size=body_size - 0.5, colour=contact_colour)

        elif section_key == "skills" and cv.skills:
            add_section_header(sections.get("skills_label", "Key Skills"))
            add_body("  •  ".join(cv.skills), space_after=4)

    # ── Save ──────────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    logger.info(f"Built DOCX for {cv.candidate.full_name or 'unknown'} "
                f"({len(cv.experience)} roles, {len(cv.education)} education entries)")
    return buf.read()
